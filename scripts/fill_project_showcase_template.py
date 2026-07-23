#!/usr/bin/env python3
"""Fill the retained project showcase template with repository-backed content."""

from __future__ import annotations

import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REFERENCE = ROOT / "tmp" / "template-whitepaper" / "reference.docx"
OUTPUT = ROOT / "output" / "docx" / "ai_agent_red_team_simulator_project_showcase.docx"

NAVY = "1A365D"
BLUE = "4F81BD"
GRAY = "4A5568"
PALE = "EAF1F6"
WHITE = "FFFFFF"


def set_run_font(run, *, size=None, bold=None, italic=None, color=None):
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def set_paragraph(paragraph, parts, *, align=None):
    """Replace paragraph text while preserving paragraph properties."""
    clear_paragraph(paragraph)
    if align is not None:
        paragraph.alignment = align
    for text, options in parts:
        run = paragraph.add_run(text)
        set_run_font(run, **options)
    return paragraph


def set_cell_text(cell, parts, *, align=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = cell.paragraphs[0]
    set_paragraph(paragraph, parts, align=align)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return paragraph


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_borders(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, attrs in edges.items():
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        for key, value in attrs.items():
            element.set(qn(f"w:{key}"), str(value))


def set_table_width(table, width_twips):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")


def add_architecture_flow(outer_cell):
    set_cell_text(outer_cell, [("", {"size": 1})])
    outer_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(outer_cell, top=160, start=160, bottom=160, end=160)

    table = outer_cell.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(1.19), Inches(1.19), Inches(1.19), Inches(1.19), Inches(1.19)]
    labels = [
        ("OPERATE", "CLI + chat"),
        ("ORCHESTRATE", "scope + monitor"),
        ("ASSESS", "static + adaptive"),
        ("EVALUATE", "detectors + severity"),
        ("REPORT", "JSON + timeline"),
    ]
    fills = [NAVY, BLUE, NAVY, BLUE, NAVY]
    for index, cell in enumerate(table.rows[0].cells):
        cell.width = widths[index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(cell, fills[index])
        set_cell_margins(cell, top=120, start=70, bottom=120, end=70)
        set_cell_borders(
            cell,
            top={"val": "single", "sz": "8", "color": WHITE},
            bottom={"val": "single", "sz": "8", "color": WHITE},
            start={"val": "single", "sz": "8", "color": WHITE},
            end={"val": "single", "sz": "8", "color": WHITE},
        )
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        title = paragraph.add_run(labels[index][0])
        set_run_font(title, size=7.4, bold=True, color=WHITE)
        subtitle = paragraph.add_run("\n" + labels[index][1])
        set_run_font(subtitle, size=6.7, color=WHITE)
    set_table_width(table, 8640)


def populate():
    if not REFERENCE.exists():
        raise FileNotFoundError(f"Template copy not found: {REFERENCE}")
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    working = ROOT / "tmp" / "template-whitepaper" / "generated-working.docx"
    shutil.copy2(REFERENCE, working)
    doc = Document(working)

    doc.core_properties.title = "AI Agent Red Team Simulator - Project Showcase"
    doc.core_properties.subject = "Enterprise AI agent security validation platform"
    doc.core_properties.author = "Vishal"
    doc.core_properties.keywords = "AI security, red team, AI agents, Kali Linux, Ollama, LangGraph"

    paragraphs = doc.paragraphs
    set_paragraph(paragraphs[0], [("AI Agent Red Team Simulator", {"size": 24, "bold": True, "color": NAVY})])
    set_paragraph(paragraphs[1], [("Authorized, local-first security testing for AI agents - from prompt behavior to tools, services, and auditable evidence.", {"size": 12, "italic": True, "color": GRAY})])

    metadata = doc.tables[0].rows[0].cells
    metadata_values = [
        ("Status: ", "Active"),
        ("Date: ", "July 2026"),
        ("Project Lead: ", "Vishal"),
        ("Quick Links: ", "GitHub Repo"),
    ]
    for cell, (label, value) in zip(metadata, metadata_values):
        set_cell_text(cell, [
            (label, {"size": 10, "bold": True, "color": GRAY}),
            (value, {"size": 10, "color": GRAY}),
        ])

    set_paragraph(paragraphs[5], [("AI agents combine probabilistic model behavior with prompts, tools, credentials, APIs, memory, and deployment infrastructure. Traditional scanners rarely test whether manipulated instructions can expose sensitive context, bypass policy, or trigger unsafe capabilities.", {"size": 10.5, "color": GRAY})])
    set_paragraph(paragraphs[7], [("The simulator provides one assessment fabric for explicit Python targets, live HTTP agents, authorized hosted endpoints, and isolated Kali workflows. Curated and adaptive probes feed a common evaluator that records severity, evidence, redacted responses, timelines, and remediation-oriented reports.", {"size": 10.5, "color": GRAY})])

    impact = doc.tables[1].cell(0, 0)
    set_cell_text(impact, [
        ("Key Impact Metric: ", {"size": 10.5, "bold": True, "color": NAVY}),
        ("47 automated tests passed, six explicitly enrolled agents discovered, and a deterministic smoke assessment completed with 6 PASS / 0 FAIL / 0 ERROR.", {"size": 10.5, "italic": True, "color": GRAY}),
    ])

    feature_parts = [
        ("Multi-path assessment: ", "Curated payloads, local adaptive probes, live HTTP discovery, and bounded Kali testing cover both model and application behavior."),
        ("Fail-closed evaluation: ", "PASS, FAIL, ERROR, and UNPARSED outcomes include severity, detector evidence, secret redaction, and automation-friendly exit behavior."),
        ("Enterprise evidence: ", "JSON, Markdown, timelines, JSONL events, and remediation guidance support release gates, triage, regression, and audit review."),
    ]
    for paragraph, (label, body) in zip(paragraphs[9:12], feature_parts):
        set_paragraph(paragraph, [
            (label, {"size": 10.3, "bold": True, "color": GRAY}),
            (body, {"size": 10.3, "color": GRAY}),
        ])

    tech = doc.tables[2]
    rows = [
        ("Orchestration", "Python CLI, natural-language assistant", "Interprets intent, scopes tests, discovers targets, and records observable assessment phases."),
        ("AI Runtime", "Ollama, LangGraph, local agents", "Keeps adaptive generation local while supporting tool-using weather and travel target services."),
        ("Security Lab", "Kali, SSH tunnels, HTTP probes", "Runs bounded recon and agent-aware tests without exposing loopback services to the LAN."),
        ("Evidence", "JSON, Markdown, JSONL", "Normalizes findings, redacts configured secrets, and produces reviewable reports and event traces."),
    ]
    for row_index, values in enumerate(rows, start=1):
        for col_index, value in enumerate(values):
            options = {"size": 9.5, "color": NAVY if col_index == 0 else GRAY}
            if col_index == 0:
                options["bold"] = True
            set_cell_text(tech.cell(row_index, col_index), [(value, options)])

    metric_parts = [
        ("Validation: ", "47 of 47 unit tests passed and the reviewed core modules compiled cleanly."),
        ("Coverage: ", "Six explicit targets were discovered; the prompt-disclosure smoke scan returned 6 PASS, 0 FAIL, 0 ERROR."),
    ]
    for paragraph, (label, body) in zip(paragraphs[15:17], metric_parts):
        set_paragraph(paragraph, [
            (label, {"size": 10.3, "bold": True, "color": GRAY}),
            (body, {"size": 10.3, "color": GRAY}),
        ])

    set_paragraph(paragraphs[17], [("Assessment Architecture", {"size": 11, "bold": True, "color": GRAY})])
    add_architecture_flow(doc.tables[3].cell(0, 0))
    set_paragraph(paragraphs[18], [("Figure 1: Authorized assessment flow from operator intent to auditable evidence.", {"size": 9, "italic": True, "color": GRAY})], align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(working)

    # python-docx rewrites unrelated package parts. Rebuild from the retained
    # reference and transplant only the intended editable parts.
    editable_parts = {"word/document.xml", "docProps/core.xml"}
    with zipfile.ZipFile(REFERENCE, "r") as source_zip, zipfile.ZipFile(working, "r") as edited_zip:
        with zipfile.ZipFile(OUTPUT, "w") as final_zip:
            for info in source_zip.infolist():
                payload = edited_zip.read(info.filename) if info.filename in editable_parts else source_zip.read(info.filename)
                final_zip.writestr(info, payload)
    working.unlink(missing_ok=True)
    return OUTPUT


if __name__ == "__main__":
    print(populate())
